"""
Ses → Metin → Word Dönüştürücü
Flask Backend: faster-whisper ile transkripsiyon, python-docx ile Word çıktısı
"""

import os
import uuid
import datetime
import threading
from pathlib import Path
import tempfile
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from docx import Document
from docx.shared import Pt

app = Flask(__name__)
CORS(app)

# Geçici dosya yükleme dizini
UPLOAD_DIR = Path(tempfile.gettempdir()) / "ses_transkript"
UPLOAD_DIR.mkdir(exist_ok=True)

# İş takibi için bellek içi sözlük ve iş parçacığı kilidi
jobs = {}
jobs_lock = threading.Lock()

# HTML sayfasını yükle
# Şablonun doğru klasörden okunmasını garanti altına almak için mutlak yol kullanıyoruz
TEMPLATE_PATH = Path(__file__).parent / "templates" / "index.html"
if TEMPLATE_PATH.exists():
    with open(TEMPLATE_PATH, encoding="utf-8") as f:
        HTML_PAGE = f.read()
else:
    HTML_PAGE = "<h1>Hata: templates/index.html dosyası bulunamadı!</h1>"


@app.route("/")
def index():
    return HTML_PAGE


@app.route("/transcribe", methods=["POST"])
def transcribe():
    if "audio" not in request.files:
        return jsonify({"error": "Ses dosyası bulunamadı"}), 400

    file = request.files["audio"]
    model = request.form.get("model", "base")
    language = request.form.get("language", "")

    if file.filename == "":
        return jsonify({"error": "Dosya seçilmedi"}), 400

    job_id = str(uuid.uuid4())
    
    with jobs_lock:
        jobs[job_id] = {
            "status": "processing", 
            "progress": "Dosya sunucuya yükleniyor ve kaydediliyor..."
        }

    # Dosya uzantısını koruyarak geçici olarak kaydet
    suffix = Path(file.filename).suffix or ".mp3"
    audio_path = UPLOAD_DIR / f"{job_id}{suffix}"
    file.save(str(audio_path))

    filename = file.filename

    def process_audio():
        try:
            with jobs_lock:
                jobs[job_id]["progress"] = "Whisper modeli belleğe yükleniyor..."
            
            # Bellek ve hız optimizasyonu için faster-whisper kütüphanesi iş parçacığı içinde içe aktarılır
            from faster_whisper import WhisperModel
            
            # İşlemci (CPU) üzerinde int8 sıkıştırma formatıyla çalıştırma (RAM tasarrufu sağlar)
            print(f"[{job_id}] Model yükleniyor: {model}")
            model_obj = WhisperModel(model, device="cpu", compute_type="int8")

            with jobs_lock:
                jobs[job_id]["progress"] = "Ses dosyası çözümleniyor (Transkripsiyon)..."
            
            opts = {}
            if language:
                opts["language"] = language

            print(f"[{job_id}] Transkripsiyon başlatıldı. Dil seçimi: {language or 'Otomatik'}")
            
            # Transkripsiyon işlemini başlat (segments bir jeneratördür)
            segments, info = model_obj.transcribe(str(audio_path), **opts)
            
            # Jeneratördeki tüm segmentleri tüketerek metni birleştiriyoruz
            text_segments = []
            for segment in segments:
                text_segments.append(segment.text)
                
            text = " ".join(text_segments).strip()
            detected_lang = info.language
            print(f"[{job_id}] Çeviri tamamlandı. Algılanan Dil: {detected_lang}")

            with jobs_lock:
                jobs[job_id].update({
                    "status": "done",
                    "text": text,
                    "language": detected_lang,
                    "filename": filename,
                })

        except Exception as e:
            print(f"[{job_id}] Kritik Hata Oluştu: {str(e)}")
            with jobs_lock:
                jobs[job_id] = {
                    "status": "error", 
                    "error": f"Dönüştürme sırasında bir hata oluştu: {str(e)}"
                }
        finally:
            # İşlem ne olursa olsun geçici ses dosyasını diskten temizle
            try:
                if audio_path.exists():
                    audio_path.unlink()
                    print(f"[{job_id}] Geçici ses dosyası silindi.")
            except Exception as clean_error:
                print(f"[{job_id}] Geçici dosya silme hatası: {str(clean_error)}")

    # İşlemi ana thread'i bloke etmemek için arka planda başlatıyoruz
    thread = threading.Thread(target=process_audio)
    thread.start()

    return jsonify({"job_id": job_id})


@app.route("/status/<job_id>")
def status(job_id):
    with jobs_lock:
        job = jobs.get(job_id)
    if not job:
        return jsonify({"error": "İşlem bulunamadı veya sunucu belleğinden silindi."}), 404
    return jsonify(job)


@app.route("/download/<job_id>")
def download(job_id):
    with jobs_lock:
        job = jobs.get(job_id)
        
    if not job or job.get("status") != "done":
        return jsonify({"error": "Transkript henüz hazır değil veya iş bulunamadı."}), 400

    text = job["text"]
    filename = Path(job.get("filename", "transkript")).stem
    detected_lang = job.get("language", "")
    
    docx_path = UPLOAD_DIR / f"{job_id}.docx"

    try:
        doc = Document()
        
        # Belge Başlığı
        baslik = doc.add_heading(f'Transkript: {filename}', 0)
        baslik.style.font.name = 'Calibri'
        
        # Bilgi ve Zaman Damgası Satırı
        now = datetime.datetime.now().strftime("%d.%m.%Y %H:%M:%S")
        bilgi_metni = f"Oluşturulma Tarihi: {now}"
        if detected_lang:
            bilgi_metni += f"    •    Algılanan Dil: {detected_lang.upper()}"
        doc.add_paragraph(bilgi_metni)
        
        doc.add_paragraph("_" * 50)  # Görsel ayırıcı çizgi
        
        # Metni paragraflara bölerek temiz bir şekilde ekle
        paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
        if not paragraphs:
            paragraphs = [text]
            
        for p_text in paragraphs:
            p = doc.add_paragraph()
            run = p.add_run(p_text)
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            
        doc.save(str(docx_path))
        
    except Exception as e:
        return jsonify({"error": f"Word belgesi oluşturulurken hata meydana geldi: {str(e)}"}), 500

    return send_file(
        str(docx_path),
        as_attachment=True,
        download_name=f"{filename}_transkript.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


if __name__ == "__main__":
    print("🎙️ Ses → Metin → Word Sunucusu başlatılıyor...")
    print("🌐 Tarayıcınızdan http://localhost:5000 adresine gidin.")
    # Port bilgisini ortam değişkenlerinden al, yoksa varsayılan olarak 5000 kullan
    port = int(os.environ.get("PORT", 5000))
    app.run(debug=False, host="0.0.0.0", port=port)
